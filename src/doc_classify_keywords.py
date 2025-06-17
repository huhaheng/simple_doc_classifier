import os
import shutil
import time
import logging
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple
import re
from datetime import datetime
import json
import sys
import pickle
from concurrent.futures import ProcessPoolExecutor, as_completed
from multiprocessing import cpu_count
import functools

__all__ = ['main']

# 需要安装: pip install python-docx python-docx2txt
try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
except ImportError as e:
    print(f"请先安装必要的库:")
    print("pip install python-docx python-docx2txt pywin32")
    print(f"缺少库: {e}")
    exit(1)

# docx2txt is optional - will be imported when needed
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

def setup_logging(base_dir: Path, resume_mode: bool = False):
    """设置轻量级日志系统"""
    # 创建日志目录
    log_dir = base_dir / "logs"
    log_dir.mkdir(parents=True, exist_ok=True)
    
    # 生成带时间戳的日志文件名
    if resume_mode:
        # 继续模式：查找最新的日志文件
        existing_logs = list(log_dir.glob("doc_classification_*.log"))
        if existing_logs:
            log_file = max(existing_logs, key=lambda x: x.stat().st_mtime)
            stats_file = log_file.with_name(log_file.stem.replace("doc_classification_", "stats_") + ".json")
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            log_file = log_dir / f"doc_classification_{timestamp}.log"
            stats_file = log_dir / f"stats_{timestamp}.json"
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        log_file = log_dir / f"doc_classification_{timestamp}.log"
        stats_file = log_dir / f"stats_{timestamp}.json"
    
    # 配置日志格式
    log_format = '%(asctime)s - %(levelname)s - %(message)s'
    date_format = '%Y-%m-%d %H:%M:%S'
    
    # 清除现有的handlers
    for handler in logging.root.handlers[:]:
        logging.root.removeHandler(handler)
    
    # 配置日志记录器
    mode = 'a' if resume_mode else 'w'  # 继续模式用追加模式
    logging.basicConfig(
        level=logging.WARNING,
        format=log_format,
        datefmt=date_format,
        handlers=[
            logging.FileHandler(log_file, encoding='utf-8', mode=mode)
        ]
    )
    
    logger = logging.getLogger(__name__)
    if resume_mode:
        logger.warning(f"=== 恢复处理任务 ===")
    else:
        logger.warning(f"=== 文档分类处理开始 ===")
    logger.warning(f"日志文件: {log_file}")
    logger.warning(f"统计文件: {stats_file}")
    
    return logger, stats_file

class ProgressManager:
    """进度管理器 - 支持断点续传"""
    def __init__(self, base_dir: Path):
        self.base_dir = base_dir
        self.progress_file = base_dir / "logs" / "processing_progress.pkl"
        self.processed_files: Set[str] = set()
        self.failed_files: Set[str] = set()
        
    def load_progress(self):
        """加载之前的进度"""
        if self.progress_file.exists():
            try:
                with open(self.progress_file, 'rb') as f:
                    data = pickle.load(f)
                    self.processed_files = data.get('processed_files', set())
                    self.failed_files = data.get('failed_files', set())
                return True
            except Exception as e:
                print(f"加载进度文件失败: {e}")
                return False
        return False
    
    def save_progress(self):
        """保存当前进度"""
        try:
            self.progress_file.parent.mkdir(parents=True, exist_ok=True)
            with open(self.progress_file, 'wb') as f:
                data = {
                    'processed_files': self.processed_files,
                    'failed_files': self.failed_files,
                    'last_update': datetime.now().isoformat()
                }
                pickle.dump(data, f)
        except Exception as e:
            print(f"保存进度文件失败: {e}")
    
    def is_processed(self, file_path: Path) -> bool:
        """检查文件是否已处理"""
        return str(file_path) in self.processed_files
    
    def mark_processed(self, file_path: Path, success: bool = True):
        """标记文件为已处理"""
        file_str = str(file_path)
        if success:
            self.processed_files.add(file_str)
            # 如果之前失败过，从失败列表中移除
            self.failed_files.discard(file_str)
        else:
            self.failed_files.add(file_str)
    
    def get_stats(self):
        """获取进度统计"""
        return {
            'processed_count': len(self.processed_files),
            'failed_count': len(self.failed_files),
            'total_attempted': len(self.processed_files) + len(self.failed_files)
        }
    
    def clear_progress(self):
        """清除进度文件"""
        if self.progress_file.exists():
            self.progress_file.unlink()
        self.processed_files.clear()
        self.failed_files.clear()

class ProcessingStats:
    """轻量级处理统计类"""
    def __init__(self, resume_stats: dict = None):
        if resume_stats:
            # 恢复之前的统计数据
            self.start_time = datetime.fromisoformat(resume_stats.get('start_time', datetime.now().isoformat()))
            self.processed_files = resume_stats.get('processed_files', 0)
            self.failed_files = resume_stats.get('failed_files', 0)
            self.empty_content_files = resume_stats.get('empty_content_files', 0)
            self.category_stats = resume_stats.get('category_stats', {})
            self.error_count = resume_stats.get('error_count', 0)
        else:
            self.start_time = datetime.now()
            self.processed_files = 0
            self.failed_files = 0
            self.empty_content_files = 0
            self.category_stats = {}
            self.error_count = 0
        
        self.end_time = None
        self.total_files = 0
        self.recent_errors = []
        
    def add_file_result(self, filename: str, category: str, success: bool, 
                       content_length: int = 0, error_msg: str = ""):
        """轻量级结果记录"""
        if success:
            self.processed_files += 1
            if category not in self.category_stats:
                self.category_stats[category] = 0
            self.category_stats[category] += 1
        else:
            self.failed_files += 1
            self.error_count += 1
            if len(self.recent_errors) >= 50:
                self.recent_errors.pop(0)
            self.recent_errors.append({
                'filename': filename,
                'error': error_msg,
                'timestamp': datetime.now().isoformat()
            })
        
        if content_length == 0:
            self.empty_content_files += 1
    
    def get_summary(self):
        """获取统计摘要"""
        self.end_time = datetime.now()
        duration = (self.end_time - self.start_time).total_seconds()
        
        return {
            'processing_summary': {
                'start_time': self.start_time.isoformat(),
                'end_time': self.end_time.isoformat(),
                'total_duration_seconds': duration,
                'total_files': self.total_files,
                'processed_files': self.processed_files,
                'failed_files': self.failed_files,
                'empty_content_files': self.empty_content_files,
                'success_rate': (self.processed_files / self.total_files * 100) if self.total_files > 0 else 0,
                'files_per_second': self.processed_files / duration if duration > 0 else 0
            },
            'category_distribution': self.category_stats,
            'recent_errors': self.recent_errors,
            # 为恢复功能保存的数据
            'category_stats': self.category_stats,
            'error_count': self.error_count
        }
    
    def save_to_file(self, stats_file: Path):
        """保存统计结果到文件"""
        summary = self.get_summary()
        with open(stats_file, 'w', encoding='utf-8') as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)

def check_resume_option(base_dir: Path) -> bool:
    """检查是否可以恢复之前的任务"""
    progress_file = base_dir / "logs" / "processing_progress.pkl"
    if not progress_file.exists():
        return False
    
    try:
        with open(progress_file, 'rb') as f:
            data = pickle.load(f)
            processed_count = len(data.get('processed_files', set()))
            failed_count = len(data.get('failed_files', set()))
            last_update = data.get('last_update', 'Unknown')
        
        print(f"\n发现未完成的处理任务:")
        print(f"  已处理文件: {processed_count}")
        print(f"  失败文件: {failed_count}")
        print(f"  最后更新: {last_update}")
        
        while True:
            choice = input("\n是否继续之前的任务? (y/n/d): ").lower().strip()
            if choice == 'y':
                return True
            elif choice == 'n':
                return False
            elif choice == 'd':
                # 删除进度文件，重新开始
                progress_file.unlink()
                print("已删除进度文件，将重新开始")
                return False
            else:
                print("请输入 y(继续), n(重新开始), 或 d(删除进度文件)")
    
    except Exception as e:
        print(f"读取进度文件失败: {e}")
        return False

def load_previous_stats(stats_file: Path) -> dict:
    """加载之前的统计数据"""
    if stats_file.exists():
        try:
            with open(stats_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return data.get('processing_summary', {})
        except:
            pass
    return {}

def get_default_config():
    """默认配置分类规则和路径"""
    config = {
        # === 路径配置 ===
        'directories': {
            'source_dir': r"C:\doc_classifier\std_docFrom20240501toNow",
            'dest_base_dir': r"C:\doc_classifier\classified_doc_v2",
            'logs_dir': "logs"
        },
        
        'categories': {
            'Benefit Summary All 01': {
                'include': ['In-Network', 'Out-of-Network', 'Benefit Effective From', 'Plan 計劃', 'HOSPITALIZATION BENEFITS', 'Outpatient Benefits', 'Other Benefits', 'Benefit Summary'],
                'exclude': ['Benefit Category', 'Coverage and Benefit Limits', 'DENTAL BENEFITS','GROUP INSURANCE', 'Hospitalization & Surgical Benefits', 'Care Plan', 'max amount', 'Policy Number', 'Effective Date', 'benefit information']
            },
            'Benefit Summary All 02': {
                'include': ['In-Network', 'Out-of-Network', 'Benefit Effective From', 'Benefit Category', 'Coverage and Benefit Limits', 'HOSPITALIZATION BENEFITS', 'Outpatient Benefits', 'Other Benefits', 'Benefit Summary'],
                'exclude': ['Plan 計劃', 'DENTAL BENEFITS','GROUP INSURANCE', 'Hospitalization & Surgical Benefits', 'Care Plan', 'max amount', 'Policy Number', 'Effective Date', 'benefit information']
            },
            'Benefit Summary All 03': {
                'include': ['Out-of-Network', 'Benefit Effective From', 'Plan 計劃', 'HOSPITALIZATION BENEFITS', 'Other Benefits', 'Benefit Summary'],
                'exclude': ['In-Network', 'Outpatient Benefits', 'Benefit Category', 'Coverage and Benefit Limits', 'DENTAL BENEFITS','GROUP INSURANCE', 'Hospitalization & Surgical Benefits', 'Care Plan', 'max amount', 'Policy Number', 'Effective Date', 'benefit information']
            },
            'Benefit Summary All 04': {
                'include': ['In-Network', 'Out-of-Network', 'Benefit Effective From', 'Plan 計劃', 'HOSPITALIZATION BENEFITS', 'Outpatient Benefits','GROUP INSURANCE', 'Other Benefits', 'Benefit Summary'],
                'exclude': ['Benefit Category', 'Coverage and Benefit Limits', 'DENTAL BENEFITS', 'Hospitalization & Surgical Benefits', 'Care Plan', 'max amount', 'Policy Number', 'Effective Date', 'benefit information']
            },
            'Benefit Summary All 05': {
                'include': ['In-Network', 'Out-of-Network', 'Benefit Effective From', 'Plan 計劃', 'HOSPITALIZATION BENEFITS', 'Outpatient Benefits', 'DENTAL BENEFITS', 'Other Benefits', 'Benefit Summary'],
                'exclude': ['Benefit Category', 'Coverage and Benefit Limits','GROUP INSURANCE', 'Hospitalization & Surgical Benefits', 'Care Plan', 'max amount', 'Policy Number', 'Effective Date', 'benefit information']
            },
            'Benefit Summary - Dental Benefits': {
                'include': ['Benefit Category', 'Coverage and Benefit Limits', 'Benefit Summary', 'Care Plan', 'Dental Benefits', 'Out-of-Network'],
                'exclude': ['In-Network', 'GROUP INSURANCE', 'Hospitalization & Surgical Benefits', 'Outpatient Benefits', 'Hospitalization Benefits', 'Maternity Benefits', 'max amount', 'Policy Number', 'Effective Date', 'benefit information', 'Benefits Schedule for Flexi Choice', 'Plan 計劃']
            },
            'Benefit Summary - Hospitalization & Surgical Benefits': {
                'include': ['In-Network', 'Out-of-Network', 'Benefit Category', 'Coverage and Benefit Limits', 'Benefit Summary', 'Care Plan','Hospitalization & Surgical Benefits'],
                'exclude': ['GROUP INSURANCE', 'Dental Benefits', 'Outpatient Benefits', 'Hospitalization Benefits', 'Maternity Benefits', 'max amount', 'Policy Number', 'Effective Date', 'benefit information', 'Benefits Schedule for Flexi Choice', 'Plan 計劃']
            },
            'Benefit Summary - Outpatient Benefits': {
                'include': ['In-Network', 'Out-of-Network', 'Benefit Category', 'Coverage and Benefit Limits', 'Benefit Summary', 'Care Plan', 'Outpatient Benefits'],
                'exclude': ['GROUP INSURANCE', 'Dental Benefits','Hospitalization & Surgical Benefits', 'Hospitalization Benefits', 'Maternity Benefits', 'max amount', 'Policy Number', 'Effective Date', 'benefit information', 'Benefits Schedule for Flexi Choice', 'Plan 計劃']
            },
            'Benefit Summary - Hospitalization Benefits': {
                'include': ['In-Network', 'Out-of-Network', 'Benefit Category', 'Coverage and Benefit Limits', 'Benefit Summary', 'Care Plan', 'Hospitalization Benefits'],
                'exclude': ['GROUP INSURANCE', 'Dental Benefits','Hospitalization & Surgical Benefits', 'Outpatient Benefits', 'Maternity Benefits', 'max amount', 'Policy Number', 'Effective Date', 'benefit information', 'Benefits Schedule for Flexi Choice', 'Plan 計劃']
            },
            'Benefit Summary - Maternity Benefits': {
                'include': ['Out-of-Network', 'Benefit Category', 'Coverage and Benefit Limits', 'Benefit Summary', 'Care Plan', 'Maternity Benefits'],
                'exclude': ['In-Network', 'GROUP INSURANCE', 'Dental Benefits','Hospitalization & Surgical Benefits', 'Outpatient Benefits', 'Hospitalization Benefits', 'max amount', 'Policy Number', 'Effective Date', 'benefit information', 'Benefits Schedule for Flexi Choice', 'Plan 計劃']
            },
            'Policy 01': {
                'include': ['benefit information', 'max amount', 'Policy Number', 'Effective Date', 'HOSPITAL AND SURGICAL'],
                'exclude': ['In-Network', 'Out-of-Network', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan', 'Benefits Schedule for Flexi Choice', 'HOSPITAL INCOME', 'NETWORK OUT-PATIENT', 'CLINICAL', 'SUPPLEMENTARY MEDICAL', 'EXTENDED MEDICAL', 'DENTAL CARE', 'NORMAL MATERNITY']
            },
            'Policy 02': {
                'include': ['benefit information', 'max amount', 'Policy Number', 'Effective Date', 'HOSPITAL AND SURGICAL', 'HOSPITAL INCOME', 'NETWORK OUT-PATIENT'],
                'exclude': ['In-Network', 'Out-of-Network', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan', 'Benefits Schedule for Flexi Choice', 'CLINICAL', 'SUPPLEMENTARY MEDICAL', 'EXTENDED MEDICAL', 'DENTAL CARE', 'NORMAL MATERNITY']
            },
            'Policy 03': {
                'include': ['benefit information', 'max amount', 'Policy Number', 'Effective Date', 'HOSPITAL AND SURGICAL', 'NETWORK OUT-PATIENT'],
                'exclude': ['In-Network', 'Out-of-Network', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan', 'Benefits Schedule for Flexi Choice', 'HOSPITAL INCOME', 'CLINICAL', 'SUPPLEMENTARY MEDICAL', 'EXTENDED MEDICAL', 'DENTAL CARE', 'NORMAL MATERNITY']
            },
            'Policy 04': {
                'include': ['benefit information', 'max amount', 'Policy Number', 'Effective Date', 'HOSPITAL AND SURGICAL', 'CLINICAL'],
                'exclude': ['In-Network', 'Out-of-Network', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan', 'Benefits Schedule for Flexi Choice', 'HOSPITAL INCOME', 'NETWORK OUT-PATIENT', 'SUPPLEMENTARY MEDICAL', 'EXTENDED MEDICAL', 'DENTAL CARE', 'NORMAL MATERNITY']
            },
            'Policy 05': {
                'include': ['benefit information', 'max amount', 'Policy Number', 'Effective Date', 'HOSPITAL AND SURGICAL', 'SUPPLEMENTARY MEDICAL'],
                'exclude': ['In-Network', 'Out-of-Network', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan', 'Benefits Schedule for Flexi Choice', 'HOSPITAL INCOME', 'NETWORK OUT-PATIENT', 'CLINICAL', 'EXTENDED MEDICAL', 'DENTAL CARE', 'NORMAL MATERNITY']
            },
            'Policy 06': {
                'include': ['benefit information', 'max amount', 'Policy Number', 'Effective Date', 'HOSPITAL AND SURGICAL', 'CLINICAL', 'EXTENDED MEDICAL'],
                'exclude': ['In-Network', 'Out-of-Network', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan', 'Benefits Schedule for Flexi Choice', 'HOSPITAL INCOME', 'NETWORK OUT-PATIENT', 'SUPPLEMENTARY MEDICAL', 'DENTAL CARE', 'NORMAL MATERNITY']
            },
            'Policy 07': {
                'include': ['benefit information', 'max amount', 'Policy Number', 'Effective Date', 'HOSPITAL AND SURGICAL', 'CLINICAL', 'SUPPLEMENTARY MEDICAL'],
                'exclude': ['In-Network', 'Out-of-Network', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan', 'Benefits Schedule for Flexi Choice', 'HOSPITAL INCOME', 'NETWORK OUT-PATIENT', 'EXTENDED MEDICAL', 'DENTAL CARE', 'NORMAL MATERNITY']
            },
            'Policy 08': {
                'include': ['benefit information', 'max amount', 'Policy Number', 'Effective Date', 'HOSPITAL AND SURGICAL', 'CLINICAL','DENTAL CARE'],
                'exclude': ['In-Network', 'Out-of-Network', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan', 'Benefits Schedule for Flexi Choice', 'HOSPITAL INCOME', 'NETWORK OUT-PATIENT', 'EXTENDED MEDICAL',  'SUPPLEMENTARY MEDICAL', 'NORMAL MATERNITY']
            },
            'Policy 09': {
                'include': ['benefit information', 'max amount', 'Policy Number', 'Effective Date', 'HOSPITAL AND SURGICAL', 'CLINICAL', 'EXTENDED MEDICAL', 'NORMAL MATERNITY'],
                'exclude': ['In-Network', 'Out-of-Network', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan', 'Benefits Schedule for Flexi Choice', 'HOSPITAL INCOME', 'NETWORK OUT-PATIENT',  'SUPPLEMENTARY MEDICAL', 'DENTAL CARE']
            },
            'Policy 10': {
                'include': ['benefit information', 'max amount', 'Policy Number', 'Effective Date', 'HOSPITAL AND SURGICAL', 'CLINICAL', 'EXTENDED MEDICAL', 'DENTAL CARE'],
                'exclude': ['In-Network', 'Out-of-Network', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan', 'Benefits Schedule for Flexi Choice', 'HOSPITAL INCOME', 'NETWORK OUT-PATIENT',  'SUPPLEMENTARY MEDICAL', 'NORMAL MATERNITY']
            },
            'Policy 11': {
                'include': ['benefit information', 'max amount', 'Policy Number', 'Effective Date', 'HOSPITAL AND SURGICAL', 'CLINICAL', 'DENTAL CARE', 'SUPPLEMENTARY MEDICAL'],
                'exclude': ['In-Network', 'Out-of-Network', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan', 'Benefits Schedule for Flexi Choice', 'HOSPITAL INCOME', 'NETWORK OUT-PATIENT', 'NORMAL MATERNITY', 'EXTENDED MEDICAL']
            },
            'Flexi Choice 01': {
                'include': ['In-Network', 'Out-of-Network', 'Benefits Schedule for Flexi Choice', 'Basic benefits'],
                'exclude': ['Core benefit', 'Optional benefits', 'Flexible Spending Arrangement', 'Adjustment Factor', 'Default plan', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan']
            },
            'Flexi Choice 02': {
                'include': ['In-Network', 'Out-of-Network', 'Benefits Schedule for Flexi Choice', 'Basic benefits', 'Optional benefits'],
                'exclude': ['Core benefit', 'Flexible Spending Arrangement', 'Adjustment Factor', 'Default plan', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan']
            },
            'Flexi Choice 03': {
                'include': ['In-Network', 'Out-of-Network', 'Benefits Schedule for Flexi Choice','Core benefit'],
                'exclude': ['Basic benefits', 'Optional benefits', 'Flexible Spending Arrangement', 'Adjustment Factor', 'Default plan', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan']
            },
            'Flexi Choice 04': {
                'include': ['In-Network', 'Out-of-Network', 'Benefits Schedule for Flexi Choice', 'Core benefit', 'Optional benefits'],
                'exclude': ['Flexible Spending Arrangement','Basic benefits', 'Adjustment Factor', 'Default plan', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan']
            },
            'Flexi Choice 05': {
                'include': ['In-Network', 'Out-of-Network', 'Benefits Schedule for Flexi Choice','Core benefit', 'Optional benefits', 'Flexible Spending Arrangement', 'Default plan'],
                'exclude': ['Basic benefits', 'Adjustment Factor', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan']
            },
            'Flexi Choice 06': {
                'include': ['In-Network', 'Out-of-Network', 'Benefits Schedule for Flexi Choice', 'Core benefit', 'Optional benefits', 'Adjustment Factor'],
                'exclude': ['Basic benefits', 'Flexible Spending Arrangement', 'Default plan', 'Benefit Category', 'Coverage and Benefit Limits', 'Plan 計劃', 'Care Plan']
            },
            'Others': {
                'include': [],
                'exclude': []
            }
        },

        # === 处理控制 ===
        'processing': {
            'delay_between_files': 0,  # 移除延时以提高性能
            'progress_interval': 50,   # 减少进度更新频率
            'auto_save_interval': 100  # 每100个文件自动保存进度
        }
    }
    return config


def read_doc_content(file_path: Path) -> tuple[str, str]:
    """读取doc/docx文件内容 - 返回(内容, 错误信息)"""
    try:
        if file_path.suffix.lower() == '.docx':
            # Optimized DOCX reading
            doc = Document(file_path)
            content_parts = []
            
            # Extract paragraph text efficiently
            content_parts.extend([p.text for p in doc.paragraphs if p.text.strip()])
            
            # Extract table text efficiently
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_text.append(cell.text.strip())
                if table_text:
                    content_parts.extend(table_text)
            
            return '\n'.join(content_parts), ""
        
        elif file_path.suffix.lower() == '.doc':
            # Try alternative methods before falling back to COM
            try:
                # Method 1: Try using antiword if available (much faster)
                import subprocess
                result = subprocess.run(['antiword', str(file_path)], 
                                      capture_output=True, text=True, timeout=30)
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip(), ""
            except (subprocess.TimeoutExpired, FileNotFoundError):
                pass
            
            # Method 2: Try using python-docx2txt (faster than COM)
            if DOCX2TXT_AVAILABLE:
                try:
                    content = docx2txt.process(str(file_path))
                    if content and content.strip():
                        return content.strip(), ""
                except Exception:
                    pass
            
            # Method 3: Fallback to COM (slow but reliable)
            word_app = None
            doc = None
            try:
                import win32com.client
                word_app = win32com.client.Dispatch("Word.Application")
                word_app.Visible = False
                word_app.DisplayAlerts = False
                doc = word_app.Documents.Open(str(file_path), ReadOnly=True)
                content = doc.Content.Text
                return content.strip() if content else "", ""
            except Exception as e:
                return "", f"读取.doc文件失败: {e}"
            finally:
                try:
                    if doc: doc.Close(SaveChanges=False)
                    if word_app: word_app.Quit()
                except: pass
                if doc: del doc
                if word_app: del word_app
        else:
            return "", "不支持的文件格式"
            
    except PackageNotFoundError:
        return "", "文件可能损坏"
    except Exception as e:
        return "", f"读取文件出错: {e}"

def classify_document(content: str, categories: Dict[str, Dict[str, List[str]]]) -> str:
    """根据内容分类文档 - 优化版本"""
    if not content or not content.strip():
        return 'Others'
    
    content_lower = content.lower()
    
    # Early exit optimization: check exclude keywords first for faster rejection
    for category_name, rules in categories.items():
        if category_name == 'Others':
            continue
            
        exclude_keywords = rules.get('exclude', [])
        
        # Fast exclude check - if any exclude keyword found, skip this category
        if exclude_keywords:
            exclude_found = False
            for keyword in exclude_keywords:
                if keyword.lower() in content_lower:
                    exclude_found = True
                    break
            if exclude_found:
                continue
        
        # Check include keywords only if no exclude keywords found
        include_keywords = rules.get('include', [])
        
        # AND logic: ALL include keywords must be present
        if include_keywords:
            all_included = True
            for keyword in include_keywords:
                if keyword.lower() not in content_lower:
                    all_included = False
                    break
            if all_included:
                return category_name
        elif not include_keywords:  # No include keywords means match by exclusion only
            return category_name
    
    return 'Others'

def create_category_folders(base_dir: Path, categories: Dict[str, Dict[str, List[str]]]) -> None:
    """创建分类文件夹"""
    for category_name in categories.keys():
        category_dir = base_dir / category_name
        category_dir.mkdir(parents=True, exist_ok=True)

def copy_file_safe(source: Path, dest_dir: Path) -> bool:
    """安全复制文件，处理重名"""
    try:
        target = dest_dir / source.name
        if target.exists():
            stem = source.stem
            suffix = source.suffix
            counter = 1
            while target.exists():
                target = dest_dir / f"{stem}_{counter}{suffix}"
                counter += 1
        
        shutil.copy2(source, target)
        return True
    except Exception:
        return False

def print_progress_bar(current: int, total: int, stats: ProcessingStats, 
                      skipped: int = 0, width: int = 50):
    """显示进度条"""
    percentage = current / total if total > 0 else 0
    elapsed = (datetime.now() - stats.start_time).total_seconds()
    
    # 计算进度条
    filled = int(width * percentage)
    bar = '█' * filled + '░' * (width - filled)
    
    # 计算预估时间
    if current > 0:
        eta_seconds = (elapsed / current) * (total - current)
        eta_str = f"{int(eta_seconds//60):02d}:{int(eta_seconds%60):02d}"
    else:
        eta_str = "--:--"
    
    # 显示进度
    status = f"\r进度: [{bar}] {current}/{total} ({percentage*100:.1f}%) | "
    status += f"成功: {stats.processed_files} | 失败: {stats.failed_files}"
    if skipped > 0:
        status += f" | 跳过: {skipped}"
    status += f" | 用时: {int(elapsed//60):02d}:{int(elapsed%60):02d} | 剩余: {eta_str}"
    
    print(status, end='', flush=True)

def process_file_worker(args: Tuple[Path, Dict, Path]) -> Tuple[str, str, bool, int, str]:
    """
    Worker function for multiprocessing file processing.
    Returns: (filename, category, success, content_length, error_msg)
    """
    file_path, categories, dest_base_dir = args
    
    try:
        # Read document content
        content, error_msg = read_doc_content(file_path)
        content_length = len(content.strip()) if content else 0
        
        if not content.strip():
            category = 'Others'
            if not error_msg:
                error_msg = "无法读取内容或文件为空"
        else:
            # Classify document
            category = classify_document(content, categories)
        
        # Copy file to appropriate folder
        dest_dir = dest_base_dir / category
        success = copy_file_safe(file_path, dest_dir)
        
        return (file_path.name, category, success, content_length, error_msg)
        
    except Exception as e:
        return (file_path.name, 'Others', False, 0, f"处理文件时出错: {str(e)}")

def get_optimal_worker_count() -> int:
    """Get optimal number of workers based on system capabilities"""
    cpu_cores = cpu_count()
    # For I/O intensive tasks, we can use more workers than CPU cores
    # But limit to reasonable number to avoid overwhelming the system
    return min(max(2, cpu_cores), 8)

def main():
    """主函数"""
    config = get_default_config()
    source_dir = Path(config['directories']['source_dir'])
    dest_base_dir = Path(config['directories']['dest_base_dir'])
    categories = config['categories']
    delay = config['processing']['delay_between_files']
    progress_interval = config['processing']['progress_interval']
    auto_save_interval = config['processing']['auto_save_interval']
    
    print("=== Doc文件内容分类工具启动 ===")
    
    # 检查源目录
    if not source_dir.exists():
        print(f"错误: 源目录不存在: {source_dir}")
        return
    
    # 创建目标目录和分类文件夹
    dest_base_dir.mkdir(parents=True, exist_ok=True)
    create_category_folders(dest_base_dir, categories)
    
    # 检查是否恢复之前的任务
    resume_mode = check_resume_option(dest_base_dir)
    
    # 设置进度管理器
    progress_mgr = ProgressManager(dest_base_dir)
    if resume_mode:
        progress_mgr.load_progress()
        prev_stats = progress_mgr.get_stats()
        print(f"恢复模式: 已处理 {prev_stats['processed_count']} 个文件")
    
    # 设置日志系统
    logger, stats_file = setup_logging(dest_base_dir, resume_mode)
    
    # 加载之前的统计数据（如果有）
    previous_stats = load_previous_stats(stats_file) if resume_mode else {}
    stats = ProcessingStats(previous_stats)
    
    # 扫描doc文件
    print("正在扫描doc文件...")
    doc_files = []
    for pattern in ['*.doc', '*.docx']:
        doc_files.extend(source_dir.rglob(pattern))
    
    if not doc_files:
        print(f"在 {source_dir} 中没有找到doc文件")
        return
    
    # 筛选出未处理的文件
    if resume_mode:
        remaining_files = [f for f in doc_files if not progress_mgr.is_processed(f)]
        skipped_count = len(doc_files) - len(remaining_files)
        print(f"总共 {len(doc_files)} 个文件，跳过已处理的 {skipped_count} 个文件")
        print(f"剩余 {len(remaining_files)} 个文件需要处理")
        doc_files = remaining_files
    
    if not doc_files:
        print("所有文件都已处理完成！")
        return
    
    stats.total_files = len(doc_files)
    doc_count = len([f for f in doc_files if f.suffix.lower() == '.doc'])
    docx_count = len([f for f in doc_files if f.suffix.lower() == '.docx'])
    
    # Determine optimal number of workers
    num_workers = get_optimal_worker_count()
    print(f"开始处理 {len(doc_files)} 个文件: {doc_count} 个.doc文件, {docx_count} 个.docx文件")
    print(f"使用 {num_workers} 个并行进程加速处理")
    print(f"按 Ctrl+C 可以安全中断并保存进度\n")
    
    # Prepare arguments for worker processes
    worker_args = [(doc_file, categories, dest_base_dir) for doc_file in doc_files]
    
    # Process files using multiprocessing
    completed_count = 0
    try:
        with ProcessPoolExecutor(max_workers=num_workers) as executor:
            # Submit all tasks
            future_to_file = {
                executor.submit(process_file_worker, args): args[0] 
                for args in worker_args
            }
            
            # Process completed tasks
            for future in as_completed(future_to_file):
                file_path = future_to_file[future]
                completed_count += 1
                
                try:
                    filename, category, success, content_length, error_msg = future.result()
                    
                    # Update statistics
                    stats.add_file_result(
                        filename=filename,
                        category=category,
                        success=success,
                        content_length=content_length,
                        error_msg=error_msg
                    )
                    
                    # Update progress manager
                    progress_mgr.mark_processed(file_path, success)
                    
                    # Log errors
                    if not success or error_msg:
                        logger.error(f"{filename}: {error_msg}")
                    
                    # Show progress
                    if completed_count % progress_interval == 0 or completed_count == len(doc_files):
                        skipped = len(progress_mgr.processed_files) - stats.processed_files if resume_mode else 0
                        print_progress_bar(completed_count, len(doc_files), stats, skipped)
                    
                    # Auto-save progress
                    if completed_count % auto_save_interval == 0:
                        progress_mgr.save_progress()
                        stats.save_to_file(stats_file)
                        
                except Exception as e:
                    logger.error(f"处理 {file_path.name} 时发生错误: {str(e)}")
                    stats.add_file_result(file_path.name, 'Others', False, 0, str(e))
                    progress_mgr.mark_processed(file_path, False)
                
    except KeyboardInterrupt:
        print(f"\n\n用户中断! 正在保存进度...")
        progress_mgr.save_progress()
        stats.save_to_file(stats_file)
        print(f"进度已保存，下次运行时可以选择继续处理")
        print(f"已处理 {stats.processed_files} 个文件")
        return
    
    # 处理完成，清理进度文件
    progress_mgr.clear_progress()
    
    # 最终进度显示
    skipped = len(progress_mgr.processed_files) - stats.processed_files if resume_mode else 0
    print_progress_bar(len(doc_files), len(doc_files), stats, skipped)
    print("\n")
    
    # 输出最终统计结果
    print("\n" + "="*80)
    print("最终统计结果:")
    print("="*80)
    
    summary = stats.get_summary()
    
    for category, count in summary['category_distribution'].items():
        print(f"{category:40s}: {count:6d} 个文件")
    
    print("-"*80)
    print(f"{'本次处理':40s}: {len(doc_files):6d} 个文件")
    print(f"{'成功处理':40s}: {stats.processed_files:6d} 个文件")
    print(f"{'处理失败':40s}: {stats.failed_files:6d} 个文件")
    print(f"{'内容为空':40s}: {stats.empty_content_files:6d} 个文件")
    print(f"{'成功率':40s}: {summary['processing_summary']['success_rate']:6.1f}%")
    print(f"{'总耗时':40s}: {summary['processing_summary']['total_duration_seconds']/60:6.1f} 分钟")
    print(f"{'平均速度':40s}: {summary['processing_summary']['files_per_second']:6.1f} 文件/秒")
    
    # 保存最终统计
    stats.save_to_file(stats_file)
    print(f"\n详细统计已保存到: {stats_file}")
    print(f"分类结果保存在: {dest_base_dir}")
    print("=== 处理完成 ===")

if __name__ == "__main__":
    main()